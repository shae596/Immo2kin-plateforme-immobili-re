#!/usr/bin/env python3
"""Génère le cahier des charges et la monographie Immo2Kin au format Word."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "word"


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n")
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x05, 0x66, 0x3B)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(14)
    r2.italic = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("\n\nPlateforme immobilière intelligente\nKinshasa — République Démocratique du Congo\n\n")
    p3.add_run("Version 1.0 — Juin 2026")

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def build_cahier_des_charges() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_cover(
        doc,
        "CAHIER DES CHARGES",
        "Immo2Kin — Plateforme immobilière intelligente",
    )

    add_heading(doc, "1. Présentation du projet")
    add_para(
        doc,
        "Immo2Kin est une plateforme web immobilière destinée au marché de Kinshasa "
        "(République Démocratique du Congo). Elle met en relation clients (locataires, "
        "acheteurs), propriétaires, agences immobilières et administrateurs autour d'un "
        "catalogue d'annonces, de réservations en ligne, de paiements, de messagerie et "
        "d'un moteur de recommandations intelligent.",
    )

    add_heading(doc, "2. Contexte et problématique", 2)
    add_bullets(
        doc,
        [
            "Fragmentation de l'offre immobilière à Kinshasa (bouche-à-oreille, réseaux sociaux, intermédiaires non structurés).",
            "Manque de transparence sur les biens (photos, prix, localisation, disponibilité).",
            "Absence d'outils numériques adaptés au contexte local (Mobile Money, communes de Kinshasa).",
            "Besoin de confiance : vérification des annonces et avis après séjour.",
        ],
    )

    add_heading(doc, "3. Objectifs", 2)
    add_heading(doc, "3.1 Objectif général", 3)
    add_para(
        doc,
        "Concevoir et déployer une solution complète de mise en relation immobilière "
        "avec gestion des annonces, réservations, paiements et recommandations personnalisées.",
    )
    add_heading(doc, "3.2 Objectifs spécifiques", 3)
    add_bullets(
        doc,
        [
            "Permettre la publication et la recherche d'annonces (location et vente).",
            "Gérer le cycle de réservation (demande, confirmation, annulation).",
            "Intégrer des moyens de paiement adaptés (Stripe, Mobile Money RDC).",
            "Offrir une messagerie temps réel entre clients et propriétaires.",
            "Renforcer la confiance via avis et vérifications administrées.",
            "Proposer des recommandations basées sur le comportement utilisateur.",
        ],
    )

    add_heading(doc, "4. Périmètre fonctionnel")
    add_table(
        doc,
        ["Phase", "Module", "Fonctionnalités livrées"],
        [
            ["0", "Infrastructure", "Monorepo, API REST, Sanctum SPA, Docker, Redis, Reverb"],
            ["1", "Authentification", "Inscription, connexion, profils, RBAC, reset mot de passe, vérification e-mail"],
            ["2", "Annonces", "CRUD biens, médias (photos/vidéos), favoris, types étendus, WhatsApp"],
            ["3", "Recherche", "Filtres avancés, carte Leaflet, géolocalisation par commune"],
            ["4", "Réservations", "Calendrier disponibilité, demandes, workflow propriétaire/client"],
            ["5", "Paiements", "Stripe, Mobile Money (Orange, Airtel, M-Pesa), webhooks"],
            ["6", "Messagerie", "Conversations par annonce, WebSocket Reverb"],
            ["7", "Confiance", "Avis post-séjour, vérifications identité/annonce, badges"],
            ["8", "Recommandations", "Scoring hybride FastAPI + fallback Laravel, annonces similaires"],
        ],
    )
    add_para(doc, "Hors périmètre retenu (Phase 9 optionnelle) : ML avancé, déploiement production complet.", bold=True)

    add_heading(doc, "5. Acteurs et rôles")
    add_table(
        doc,
        ["Rôle", "Description", "Droits principaux"],
        [
            ["Client", "Locataire ou acheteur", "Recherche, favoris, réservations, paiements, messages, avis"],
            ["Propriétaire", "Propriétaire de biens", "CRUD annonces, gestion réservations, messagerie, vérification"],
            ["Agence", "Agence immobilière", "Mêmes droits que propriétaire sur plusieurs biens"],
            ["Administrateur", "Équipe plateforme", "Back-office : utilisateurs, annonces, paiements, vérifications"],
        ],
    )

    add_heading(doc, "6. Exigences fonctionnelles détaillées")
    sections = [
        (
            "6.1 Gestion des comptes",
            [
                "Inscription avec choix du rôle (client, propriétaire, agence).",
                "Authentification sécurisée par cookies Sanctum (SPA).",
                "Profil : nom, téléphone, ville, commune, avatar, biographie.",
                "Réinitialisation et vérification d'adresse e-mail.",
            ],
        ),
        (
            "6.2 Annonces immobilières",
            [
                "Types : appartement, maison, villa, studio, bureau, commerce.",
                "Transaction : location (rent) ou vente (sale).",
                "Statuts : brouillon, publié, archivé.",
                "Caractéristiques : prix (USD), chambres, salles de bain, surface, équipements.",
                "Upload photos (JPG, PNG, WebP, max 10 Mo) et vidéos.",
                "Suppression de photos existantes en mode édition.",
            ],
        ),
        (
            "6.3 Recherche et découverte",
            [
                "Liste paginée avec filtres (ville, commune, type, prix, surface).",
                "Tri : plus récent, prix croissant/décroissant, surface.",
                "Carte interactive (Leaflet) avec marqueurs géolocalisés.",
                "Section « Recommandé pour vous » et annonces similaires.",
            ],
        ),
        (
            "6.4 Réservations",
            [
                "Réservation uniquement sur annonces en location publiées.",
                "Vérification des chevauchements de dates (pending + confirmed).",
                "Workflow : pending → confirmed / rejected / cancelled.",
                "Calendrier de disponibilité exposé via API.",
            ],
        ),
        (
            "6.5 Paiements",
            [
                "Stripe (PaymentIntent, webhook).",
                "Mobile Money RDC : Orange, Airtel, M-Pesa (simulation en dev).",
                "Association paiement ↔ réservation, horodatage paid_at.",
            ],
        ),
        (
            "6.6 Messagerie",
            [
                "Conversation unique client ↔ propriétaire par annonce.",
                "Messages en temps réel via Laravel Reverb.",
                "Contact WhatsApp depuis la fiche annonce.",
            ],
        ),
        (
            "6.7 Avis et vérifications",
            [
                "Avis 1–5 étoiles après séjour confirmé et terminé.",
                "Demande de vérification identité ou annonce (validation admin).",
                "Badges « Vérifié » sur profils et annonces.",
            ],
        ),
    ]
    for title, items in sections:
        add_heading(doc, title, 2)
        add_bullets(doc, items)

    add_heading(doc, "7. Exigences non fonctionnelles")
    add_table(
        doc,
        ["Catégorie", "Exigence"],
        [
            ["Performance", "Pagination API (12 éléments/page), cache Redis, queues asynchrones"],
            ["Sécurité", "CSRF, CORS, RBAC Spatie, policies Laravel, throttle login/uploads"],
            ["Disponibilité", "Fallback recommandations si service IA indisponible"],
            ["Maintenabilité", "Clean architecture : Controllers, Services, Repositories, Resources"],
            ["Testabilité", "69 tests Feature PHPUnit backend"],
            ["UX", "Interface responsive React + Tailwind CSS 4"],
            ["Localisation", "Kinshasa : communes, coordonnées GPS, Mobile Money RDC"],
        ],
    )

    add_heading(doc, "8. Architecture technique")
    add_table(
        doc,
        ["Couche", "Technologie"],
        [
            ["API", "Laravel 12, PHP 8.2+, Sanctum, API REST /api/v1"],
            ["Frontend", "React 19, TypeScript, Vite 8, Tailwind CSS 4, Zustand"],
            ["Temps réel", "Laravel Reverb, Laravel Echo, Pusher protocol"],
            ["Base de données", "MySQL 8.4 (SQLite en mode fallback dev)"],
            ["Cache / files", "Redis 7"],
            ["IA", "FastAPI (Python), microservice recommandations port 8001"],
            ["Infra dev", "Docker Compose (MySQL, Redis, AI service)"],
        ],
    )

    add_heading(doc, "9. Contraintes et hypothèses")
    add_bullets(
        doc,
        [
            "Devise principale : USD.",
            "Zone géographique cible : Kinshasa (extensible à d'autres villes).",
            "Navigateurs modernes (Chrome, Firefox, Edge).",
            "Connexion Internet requise pour paiements et messagerie temps réel.",
            "Upload photos : limite PHP 12 Mo en environnement de développement configuré.",
        ],
    )

    add_heading(doc, "10. Livrables")
    add_bullets(
        doc,
        [
            "Code source monorepo (backend, frontend, ai-service).",
            "Documentation technique (architecture, API, base de données).",
            "Scripts de démarrage et configuration (.env.example).",
            "Données de démonstration (seeders : 8 annonces, comptes test).",
            "Suite de tests automatisés (69 tests Feature).",
        ],
    )

    add_heading(doc, "11. Critères d'acceptation")
    add_bullets(
        doc,
        [
            "Un client peut rechercher, réserver et payer une annonce en location.",
            "Un propriétaire peut publier, modifier et gérer ses annonces et réservations.",
            "L'administrateur peut modérer utilisateurs, annonces et vérifications.",
            "Les recommandations s'affichent sur l'accueil et les fiches annonces.",
            "Tous les tests Feature passent (69/69).",
        ],
    )

    add_heading(doc, "12. Planning de réalisation (phases livrées)")
    add_table(
        doc,
        ["Phase", "Intitulé", "Statut"],
        [
            ["0", "Bootstrap & infrastructure", "Terminée"],
            ["1", "Authentification & profils", "Terminée"],
            ["2", "Annonces & favoris", "Terminée"],
            ["3", "Recherche & carte", "Terminée"],
            ["4", "Réservations", "Terminée"],
            ["5", "Paiements", "Terminée"],
            ["6", "Messagerie temps réel", "Terminée"],
            ["7", "Avis & vérifications", "Terminée"],
            ["8", "Recommandations IA", "Terminée"],
            ["9", "ML avancé & production", "Optionnelle / hors périmètre"],
        ],
    )

    return doc


def build_monographie() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_cover(
        doc,
        "MONOGRAPHIE",
        "Conception et réalisation d'Immo2Kin",
    )

    add_heading(doc, "Résumé")
    add_para(
        doc,
        "Immo2Kin est une plateforme immobilière web intelligente conçue pour le marché "
        "kinshais. Ce document présente le contexte, l'analyse des besoins, les choix "
        "architecturaux, l'implémentation technique et les résultats obtenus. La solution "
        "repose sur une architecture monorepo combinant une API Laravel 12, une SPA React 19 "
        "et un microservice FastAPI de recommandations. Les phases 0 à 8 couvrent "
        "l'ensemble du cycle métier : de l'authentification aux paiements Mobile Money, "
        "en passant par la messagerie temps réel et la confiance utilisateur.",
    )
    add_para(doc, "Mots-clés : immobilier, Kinshasa, Laravel, React, recommandations, Mobile Money, RDC.")

    add_heading(doc, "Table des matières (structure)")
    add_bullets(
        doc,
        [
            "Chapitre I — Introduction et problématique",
            "Chapitre II — Analyse et spécifications",
            "Chapitre III — Conception et architecture",
            "Chapitre IV — Réalisation technique",
            "Chapitre V — Tests et validation",
            "Chapitre VI — Résultats et perspectives",
            "Conclusion",
            "Bibliographie et webographie",
        ],
    )
    doc.add_page_break()

    add_heading(doc, "Chapitre I — Introduction et problématique")
    add_heading(doc, "1.1 Contexte", 2)
    add_para(
        doc,
        "Le secteur immobilier à Kinshasa connaît une forte demande pour des logements, "
        "bureaux et locaux commerciaux. Cependant, la digitalisation reste limitée : les "
        "transactions passent souvent par des intermédiaires informels, sans traçabilité ni "
        "garanties. Immo2Kin répond à ce besoin en centralisant l'offre et en digitalisant "
        "les interactions entre les parties prenantes.",
    )
    add_heading(doc, "1.2 Problématique", 2)
    add_para(
        doc,
        "Comment concevoir une plateforme immobilière complète, adaptée au contexte "
        "congolais, intégrant paiements locaux, géolocalisation par commune et "
        "recommandations intelligentes, tout en garantissant sécurité et évolutivité ?",
    )
    add_heading(doc, "1.3 Objectifs de la monographie", 2)
    add_bullets(
        doc,
        [
            "Documenter les choix techniques et méthodologiques du projet Immo2Kin.",
            "Présenter l'architecture logicielle et le modèle de données.",
            "Analyser les modules livrés et leur contribution à la valeur métier.",
            "Identifier les limites actuelles et les évolutions possibles.",
        ],
    )

    add_heading(doc, "Chapitre II — Analyse et spécifications")
    add_heading(doc, "2.1 Étude des besoins", 2)
    add_para(
        doc,
        "L'analyse a identifié quatre profils utilisateurs : client, propriétaire, agence "
        "et administrateur. Les besoins fonctionnels ont été organisés en neuf phases "
        "incrémentales, de l'infrastructure de base (Phase 0) aux recommandations IA (Phase 8).",
    )
    add_heading(doc, "2.2 Modélisation des données", 2)
    add_para(
        doc,
        "Le schéma relationnel MySQL comprend les entités principales suivantes :",
    )
    add_bullets(
        doc,
        [
            "users — comptes avec profil étendu et rôles Spatie Permission",
            "properties — annonces (statut, type, listing_type, géolocalisation)",
            "property_images / property_videos — médias",
            "amenities — équipements liés aux biens",
            "favorites — favoris utilisateur",
            "reservations — demandes de location avec dates et statuts",
            "payments — transactions Stripe et Mobile Money",
            "conversations / messages — messagerie",
            "reviews — avis clients",
            "verifications — demandes de certification",
            "recommendation_events — signaux pour le moteur de recommandations",
        ],
    )
    add_heading(doc, "2.3 Cas d'utilisation principaux", 2)
    add_table(
        doc,
        ["Cas d'utilisation", "Acteur", "Description"],
        [
            ["Publier une annonce", "Propriétaire", "Création, photos, publication ou brouillon"],
            ["Réserver un bien", "Client", "Sélection dates, message, demande au propriétaire"],
            ["Payer une réservation", "Client", "Stripe ou Mobile Money"],
            ["Échanger avec le propriétaire", "Client", "Messagerie temps réel ou WhatsApp"],
            ["Modérer une vérification", "Admin", "Approuver ou rejeter une demande"],
            ["Consulter recommandations", "Client", "Accueil personnalisé et biens similaires"],
        ],
    )

    add_heading(doc, "Chapitre III — Conception et architecture")
    add_heading(doc, "3.1 Architecture globale", 2)
    add_para(
        doc,
        "Immo2Kin adopte une architecture en trois couches applicatives au sein d'un "
        "monorepo :",
    )
    add_bullets(
        doc,
        [
            "backend/ — API REST Laravel (logique métier, persistance, auth)",
            "frontend/ — SPA React consommant l'API via Axios (proxy Vite en dev)",
            "ai-service/ — Microservice FastAPI pour le scoring de recommandations",
        ],
    )
    add_heading(doc, "3.2 Clean architecture (backend)", 2)
    add_para(
        doc,
        "Le backend Laravel suit une séparation en couches :",
    )
    add_bullets(
        doc,
        [
            "Controllers (minces) — routage HTTP, délégation aux services",
            "Requests — validation des entrées",
            "Resources — sérialisation JSON des réponses",
            "Services — règles métier (PropertyService, ReservationService, etc.)",
            "Repositories — requêtes base de données",
            "Policies — autorisation fine par ressource",
        ],
    )
    add_heading(doc, "3.3 Authentification et sécurité", 2)
    add_para(
        doc,
        "L'authentification utilise Laravel Sanctum en mode SPA : cookie de session, "
        "protection CSRF, CORS configuré pour le frontend. Le RBAC s'appuie sur Spatie "
        "Laravel Permission (rôles client, proprietaire, agence, admin).",
    )
    add_heading(doc, "3.4 Temps réel", 2)
    add_para(
        doc,
        "La messagerie s'appuie sur Laravel Reverb (WebSocket) et Laravel Echo côté "
        "frontend. Les messages sont diffusés sur le canal conversation.{id}.",
    )
    add_heading(doc, "3.5 Moteur de recommandations", 2)
    add_para(
        doc,
        "Le module de recommandations combine :",
    )
    add_bullets(
        doc,
        [
            "Collecte d'événements (vues, favoris, réservations, avis, recherches)",
            "Appel au microservice FastAPI (POST /api/v1/recommendations/rank)",
            "Fallback local Laravel (RecommendationScorer) si l'IA est indisponible",
            "Affichage : section accueil et annonces similaires sur fiche produit",
        ],
    )

    add_heading(doc, "Chapitre IV — Réalisation technique")
    add_heading(doc, "4.1 Stack technologique", 2)
    add_table(
        doc,
        ["Composant", "Version / outil", "Rôle"],
        [
            ["PHP", "8.2+", "Runtime backend"],
            ["Laravel", "12", "Framework API"],
            ["React", "19", "Interface utilisateur"],
            ["TypeScript", "6", "Typage frontend"],
            ["Vite", "8", "Build et serveur de dev"],
            ["Tailwind CSS", "4", "Styles utilitaires"],
            ["MySQL", "8.4", "Persistance relationnelle"],
            ["Redis", "7", "Cache, sessions, queues"],
            ["FastAPI", "Python 3.12", "Microservice IA"],
            ["Leaflet", "1.9", "Cartographie"],
            ["Stripe", "API", "Paiements carte"],
        ],
    )

    add_heading(doc, "4.2 Modules frontend", 2)
    add_bullets(
        doc,
        [
            "Pages publiques : accueil, catalogue, carte, détail annonce",
            "Espace client : favoris, réservations, messages, dashboard",
            "Espace propriétaire : mes annonces, édition, gestion réservations, vérification",
            "Back-office admin : utilisateurs, annonces, réservations, paiements, vérifications",
            "Composants : ReservationForm, PaymentDialog, PropertyReviewsSection, SimilarProperties",
        ],
    )

    add_heading(doc, "4.3 API REST", 2)
    add_para(
        doc,
        "L'API est versionnée sous le préfixe /api/v1. Elle expose plus de 50 endpoints "
        "couvrant l'authentification, les annonces, les médias, les réservations, les "
        "paiements, la messagerie, les avis, les vérifications et les recommandations. "
        "La documentation complète est disponible dans docs/api.md.",
    )

    add_heading(doc, "4.4 Environnement de développement", 2)
    add_bullets(
        doc,
        [
            "Scripts PowerShell : setup.ps1, start-dev.ps1, serve-backend.ps1",
            "Serveur API : PHP built-in server (port 8000) avec limites upload 12 Mo",
            "Frontend Vite (port 5173) avec proxy /api et /sanctum",
            "Docker Compose : MySQL (3307), Redis (6379), AI service (8001)",
        ],
    )

    add_heading(doc, "Chapitre V — Tests et validation")
    add_heading(doc, "5.1 Stratégie de tests", 2)
    add_para(
        doc,
        "Le backend dispose de 69 tests Feature PHPUnit couvrant les domaines critiques :",
    )
    add_table(
        doc,
        ["Suite de tests", "Domaine"],
        [
            ["AuthTest", "Inscription, connexion, profil"],
            ["PropertyTest", "CRUD annonces, upload images"],
            ["PropertySearchTest", "Recherche et filtres"],
            ["ReservationTest", "Réservations et disponibilité"],
            ["PaymentTest", "Stripe et Mobile Money"],
            ["ConversationTest", "Messagerie"],
            ["ReviewTest / VerificationTest", "Avis et vérifications"],
            ["RecommendationTest", "Recommandations"],
            ["AdminPanelTest", "Back-office"],
        ],
    )
    add_heading(doc, "5.2 Validation fonctionnelle", 2)
    add_para(
        doc,
        "Des comptes de démonstration permettent de valider les parcours utilisateur : "
        "client@immo.local, proprietaire@immo.local, admin@immo.local (mot de passe : password). "
        "Huit annonces de démo couvrent les communes de Kinshasa (Gombe, Limete, Ngaliema, etc.).",
    )

    add_heading(doc, "Chapitre VI — Résultats et perspectives")
    add_heading(doc, "6.1 Résultats obtenus", 2)
    add_bullets(
        doc,
        [
            "Plateforme fonctionnelle couvrant 8 phases métier sur 9 prévues.",
            "Architecture modulaire et testée (69/69 tests).",
            "Interface moderne responsive en français.",
            "Intégration paiements locaux (Mobile Money) et internationaux (Stripe).",
            "Moteur de recommandations hybride avec résilience (fallback).",
        ],
    )
    add_heading(doc, "6.2 Limites actuelles", 2)
    add_bullets(
        doc,
        [
            "Déploiement production non finalisé (Phase 9 hors périmètre).",
            "Pas de Nginx/Apache configuré — serveurs de développement uniquement.",
            "Mobile Money en mode simulation en environnement de dev.",
            "Géolocalisation par commune (coordonnées approximatives, pas d'adresse exacte GPS).",
        ],
    )
    add_heading(doc, "6.3 Perspectives d'évolution", 2)
    add_bullets(
        doc,
        [
            "Phase 9 : ML avancé, déploiement production (Nginx, CI/CD).",
            "Application mobile (React Native ou PWA).",
            "Extension à d'autres villes de la RDC.",
            "Intégration API Mobile Money réelles (Orange Money, Airtel Money).",
            "Notifications push et alertes prix.",
        ],
    )

    add_heading(doc, "Conclusion")
    add_para(
        doc,
        "Immo2Kin démontre la faisabilité d'une plateforme immobilière complète adaptée "
        "au contexte kinshais. L'architecture monorepo, la séparation des responsabilités "
        "et l'approche incrémentale par phases ont permis de livrer un produit cohérent "
        "couvrant l'ensemble de la chaîne de valeur immobilière numérique. Les travaux "
        "futurs porteront sur le déploiement en production et l'enrichissement du moteur "
        "d'intelligence artificielle.",
    )

    add_heading(doc, "Bibliographie et webographie")
    add_bullets(
        doc,
        [
            "Documentation Laravel 12 — https://laravel.com/docs",
            "Documentation React 19 — https://react.dev",
            "FastAPI — https://fastapi.tiangolo.com",
            "Laravel Sanctum SPA Authentication — https://laravel.com/docs/sanctum",
            "Spatie Laravel Permission — https://spatie.be/docs/laravel-permission",
            "Leaflet — https://leafletjs.com",
            "Stripe API — https://stripe.com/docs/api",
        ],
    )

    return doc


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)

    cdc_path = OUT / "Immo2Kin_Cahier_des_charges.docx"
    mono_path = OUT / "Immo2Kin_Monographie.docx"

    build_cahier_des_charges().save(str(cdc_path))
    build_monographie().save(str(mono_path))

    print(f"Généré : {cdc_path}")
    print(f"Généré : {mono_path}")


if __name__ == "__main__":
    main()
